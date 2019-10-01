from docx import image as docImg
import re
import random

import sys
from spider import request_util

'''试题爬取基类'''
class ExamLib:
    
    '''图片后缀正则'''
    img_suffix_pattern = re.compile('\.\w+$')
    html_suffix_pattern = re.compile('\.html$')
    '''html页面编码'''
    encoding = 'gb2312'

    '''将div内容添加到document'''
    def add_paragraph(self, document, div, is_depth=False):
        if document == None or div == None:
            return
        for p in div.findAll('p'):
            document.add_paragraph(p.text)
            img = p.find('img')
            if img != None:
                img_url = img['src']
                img_type = self.img_suffix_pattern.search(img_url).group(0)
                img_path = self.file_dir+'/'+str(random.randint(1, 1000000))+img_type                               
                if request_util.download_img(img_path, img_url):
                    try:
                        document.add_picture(img_path)
                    except docImg.exceptions.UnexpectedEndOfFileError:
                        os.remove(img_path)
            # 是否不递归爬取页面的a标签
            if not is_depth:
                for a in p.findAll('a'):
                    if self.html_suffix_pattern.match(a['href']):
                        self.request_exam_page(a['href'], document, True)

    def request_exam_page(self, url, word, is_depth):
        pass
