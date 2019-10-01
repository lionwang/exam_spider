from docx import Document
from bs4 import BeautifulSoup

import os

import sys
from spider import request_util
from spider.exam_lib.exam_lib import ExamLib

'''
中公公考题目爬取类
'''
class OffcnExamLib(ExamLib):

    base_url = 'http://www.offcn.com/gjgwy/shiti/'
    type_dict = {'行测': '2966', '申论': '2968', '面试': '2970'}

    '''构造函数'''
    def __init__(self, file_dir='./'):
        self.file_dir = file_dir

    '''请求页面并保存到文件内'''
    def request(self, file_name, exam_type='行测', from_page=1, to_page=5):
        word = Document()
        type_value = self.type_dict.get(exam_type)
        if type_value != None:
            type_url = self.base_url + type_value + '/'
            print('---------开始下载试题内容----------')
            for i in range(from_page, to_page+1):
                url = type_url
                if i > 1:
                    url += str(i)+'.html'
                self.request_index_page(url, word)
            word.save(self.file_dir+'/'+file_name)
            print('\n---------下载试题完成，到'+self.file_dir+'目录查看-----------')
        else:
            print('------输入试题类型不合法-----')

    '''解析试题索引页面'''
    def request_index_page(self, url, word):
        html = request_util.get_html(url, self.encoding)
        if html != None:
            doc = BeautifulSoup(html, features='html5lib')
            index_ul = doc.find('ul', class_='lh_newBobotm02')
            if index_ul != None:
                for li in index_ul.findAll('li'):
                    if li != None:
                        a = li.findAll('a')
                        if len(a) > 1:
                            self.request_exam_page(a[1]['href'], word)
                            print('>', end='', flush=True)

    '''解析试题页面'''
    def request_exam_page(self, url, word, is_depth=False):
        html = request_util.get_html(url, self.encoding)
        if html != None:
            doc = BeautifulSoup(html, features='html5lib')        
            self.add_paragraph(word, doc.find('div', class_='offcn_shocont'), is_depth)

    '''将div内容添加到document'''
    '''
    def add_paragraph(self, document, div, is_depth=False):
        if document == None or div == None:
            return
        for p in div.findAll('p'):
            document.add_paragraph(p.text)
            img = p.find('img')
            if img != None:
                img_url = img['src']
                img_type = self.img_suffix_pattern.search(img_url).group(0)
                img_path = self.file_dir+'/'+str(random.randint(1, 1000000))+img_type                               
                if request_util.download_img(img_path, img_url):
                    try:
                        document.add_picture(img_path)
                    except docImg.exceptions.UnexpectedEndOfFileError:
                        os.remove(img_path)
            # 是否不递归爬取页面的a标签
            if not is_depth:
                for a in p.findAll('a'):
                    if self.html_suffix_pattern.match(a['href']):
                        self.request_exam_page(a['href'], document, True)
        '''


